from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some interesting text probably')
p.add_run(' bold text ').bold = True
p.add_run(' and also ')
p.add_run(' italic text.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='IntenseQuote')

document.add_paragraph(
	'first item in unordered list', style='ListBullet'
)

document.add_paragraph(
	'first item in ordered list', style='ListNumber'
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'

#This is how to populate from a list
'''for item in recordset:
	row_cells = table.add_row().cells
	row_cells[0].text = str(item.qty)
	row_cells[1].text = str(item.id)
	row_cells[2].text = str(item.desc)'''

document.add_page_break()

document.save('demo.docx')
