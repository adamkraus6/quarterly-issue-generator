#!/usr/bin/python

# =============================================================================
# This provides functionality for reading the log .csv file into a list! Nice!
# =============================================================================

import csv
import Tkinter as tk
import tkFileDialog
import zipfile
import os
import collections
from docx import Document
from datetime import date


def readcsv():
	Log = collections.namedtuple("Log", ["date", "time", "psa"])
	logs = []

	root = tk.Tk()
	root.withdraw()

	file_path = tkFileDialog.askopenfilename()
	zip_ref = zipfile.ZipFile(file_path, 'r')
	zip_ref.extractall('logs/extracted/')
	zip_ref.close()

	path = 'logs/extracted/' + os.path.splitext(os.path.basename(file_path))[0]+'/'

	for filename in os.listdir(path):
		print(filename)
		if not filename.startswith('.'):
			with open(path+filename) as csvfile:
				reader = csv.reader(csvfile)
				for row in reader:
					entry = Log(row[0], row[1], row[2])
					logs.append(entry)
	return logs

def generatedocx(logs, quarter):
	# This function will generate the .docx file for the report, given the list
	# of logs in the form of the named tuple.

	#create the document
	report = Document()

	logs.sort(key=lambda tup: tup[2])

	testlogs = {}
	for log in logs:
		if log.psa not in testlogs:
			testlogs[log.psa] = {}
                if log.date not in testlogs[log.psa]:
			testlogs[log.psa][log.date] = []
		testlogs[log.psa][log.date].append(log.time)

	#Generate the title
	if date.today().month == 1:
		report.add_heading('KTEQ-FM Quarterly Issues Report Q' + str(quarter) + ' '+ str(date.today().year-1), level=0)
	else:
		report.add_heading('KTEQ-FM Quarterly Issues Report Q' + str(quarter) + ' '+ str(date.today().year), level=0)

	report.add_paragraph('This document is the quarterly Community Issues Report for KTEQ-FM. It details a number of community issues discussed during programming throughout the quarter, and lists public service announcements that support these issues. This list contains all of the public service announcements played on air by live DJs. For a complete list, including automated public service announcements, contact KTEQ-FM management at kteq@mines.sdsmt.edu')

	table = report.add_table(rows=1, cols=3)
	hdr = table.rows[0].cells
	hdr[0].text = 'Date Played'
	hdr[1].text = 'Time Played'
	hdr[2].text = 'PSA Title'

	for entry in logs:
		row = table.add_row().cells
		row[0].text = str(entry.date)
		row[1].text = str(entry.time)
		row[2].text = entry.psa

	'''for psa in testlogs:
		report.add_heading(psa, level=1)
		table = report.add_table(rows=len(testlogs[psa]), cols=2)
		for psadate in testlogs[psa]:
			row = table.add_row().cells
			row[0].text = str(psadate)
			row[1].text = str(testlogs[psa][psadate])'''

	reportdir = 'reports/' + str(date.today().year)
	if not os.path.exists(reportdir):
		os.makedirs(reportdir)

	outputfile = reportdir + '/Q' + str(quarter) + '_Issues.docx'
	report.save(outputfile)
